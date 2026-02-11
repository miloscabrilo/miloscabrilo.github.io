#!/usr/bin/env python3
"""
Generate JPR_2026.docx - Jedinstvena prijava za registraciju poreskih obveznika,
obveznika doprinosa i osiguranika (OBRAZAC JPR)

This script creates an editable DOCX file matching the layout of the original PDF.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right, each with val, sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = parse_xml(f'<w:{edge} {nsdecls("w")}/>')
            tcBorders.append(element)
        for attr_name, attr_val in attrs.items():
            element.set(qn(f'w:{attr_name}'), str(attr_val))


def set_cell_width(cell, width_cm):
    """Set cell width."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")}/>')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # cm to twips
    tcW.set(qn('w:type'), 'dxa')


def set_row_height(row, height_cm):
    """Set row height."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
        tr.insert(0, trPr)
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")}/>')
        trPr.append(trHeight)
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')


def add_checkbox(paragraph, label="", checked=False):
    """Add a checkbox character followed by label text."""
    run = paragraph.add_run("☐ " if not checked else "☒ ")
    run.font.size = Pt(11)
    if label:
        run2 = paragraph.add_run(label)
        run2.font.size = Pt(10)


def make_paragraph(doc_or_cell, text="", bold=False, size=10, alignment=None, space_after=0, space_before=0, italic=False):
    """Add a formatted paragraph."""
    if hasattr(doc_or_cell, 'add_paragraph'):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def clear_cell(cell):
    """Clear all paragraphs from a cell."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""


def cell_text(cell, text, size=9, bold=False, alignment=None):
    """Set cell text with formatting."""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    return p


def set_narrow_margins(section):
    """Set narrow page margins."""
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def add_input_cells_row(table, label, num_cells, label_width=4.0, cell_width=0.65):
    """Add a row with label and individual input cells for characters."""
    row = table.add_row()
    cell_text(row.cells[0], label, size=9, bold=True)
    set_cell_width(row.cells[0], label_width)
    return row


def create_field_table(doc, rows_data, label_width=4.5, field_width=12.0):
    """Create a simple label-value table for form fields.
    rows_data: list of (label, default_value) tuples.
    """
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(rows_data):
        cell_text(table.rows[i].cells[0], label, size=9, bold=True)
        set_cell_width(table.rows[i].cells[0], label_width)
        cell_text(table.rows[i].cells[1], value, size=9)
        set_cell_width(table.rows[i].cells[1], field_width)
        # Add bottom border to the value cell
        set_cell_border(table.rows[i].cells[1],
                        bottom={'val': 'single', 'sz': '4', 'color': '000000'})
    return table


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ============================================================
# MAIN DOCUMENT CREATION
# ============================================================

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

# Set up first section margins
section = doc.sections[0]
set_narrow_margins(section)
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)

# ============================================================
# PAGE 1: Main Form
# ============================================================

# Header - right aligned "OBRAZAC JPR"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('„OBRAZAC JPR')
run.font.size = Pt(11)
run.bold = True

# CRNA GORA / PORESKA UPRAVA
p = doc.add_paragraph()
run = p.add_run('CRNA GORA\nPORESKA UPRAVA')
run.font.size = Pt(10)
run.bold = True
p.paragraph_format.space_after = Pt(6)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Jedinstvena prijava za registraciju poreskih obveznika,\nobveznika doprinosa i osiguranika')
run.font.size = Pt(14)
run.bold = True
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)

# Tip dokumenta section
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tip dokumenta')
run.font.size = Pt(10)
run.bold = True
p.paragraph_format.space_after = Pt(4)

# Checkboxes for document type - using a table for alignment
tip_table = doc.add_table(rows=5, cols=2)
tip_table.alignment = WD_TABLE_ALIGNMENT.CENTER
checkboxes_data = [
    'Registracija - Upis u registar',
    'Izmjena podataka',
    'Prestanak registracije - Brisanje iz registra',
    'Prijava/odjava osiguranika',
    'Evidencija ovlašćenih lica'
]
for i, cb_text in enumerate(checkboxes_data):
    cell_text(tip_table.rows[i].cells[0], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_width(tip_table.rows[i].cells[0], 1.0)
    cell_text(tip_table.rows[i].cells[1], cb_text, size=10)
    set_cell_width(tip_table.rows[i].cells[1], 8.0)

# Remove borders from tip table
for row in tip_table.rows:
    for cell in row.cells:
        set_cell_border(cell,
                        top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                        bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                        left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                        right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# Note about X
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Upisati «X» u jedno od navedenih polja)')
run.font.size = Pt(9)
run.italic = True
p.paragraph_format.space_after = Pt(8)

# ---- Section 1: Podaci o podnosiocu ----
p = doc.add_paragraph()
run = p.add_run('1. Podaci o podnosiocu')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_after = Pt(4)

# 1.1 Tip lica
tip_lica_table = doc.add_table(rows=1, cols=5)
tip_lica_table.alignment = WD_TABLE_ALIGNMENT.LEFT
cell_text(tip_lica_table.rows[0].cells[0], '1.1 Tip lica', size=9, bold=True)
set_cell_width(tip_lica_table.rows[0].cells[0], 3.5)
cell_text(tip_lica_table.rows[0].cells[1], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_width(tip_lica_table.rows[0].cells[1], 0.8)
cell_text(tip_lica_table.rows[0].cells[2], 'Pravno lice', size=9, bold=True)
set_cell_width(tip_lica_table.rows[0].cells[2], 4.0)
cell_text(tip_lica_table.rows[0].cells[3], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_width(tip_lica_table.rows[0].cells[3], 0.8)
cell_text(tip_lica_table.rows[0].cells[4], 'Fizičko lice', size=9, bold=True)
set_cell_width(tip_lica_table.rows[0].cells[4], 4.0)
for cell in tip_lica_table.rows[0].cells:
    set_cell_border(cell,
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 1.2 Matični broj/JMB - individual cells
jmb_table = doc.add_table(rows=1, cols=14)
jmb_table.alignment = WD_TABLE_ALIGNMENT.LEFT

# Add label before cells
p_jmb = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()

# Actually let's use a different approach - label + cells table
jmb_label_table = doc.add_table(rows=1, cols=2)
cell_text(jmb_label_table.rows[0].cells[0], '1.2 Matični broj/JMB', size=9, bold=True)
set_cell_width(jmb_label_table.rows[0].cells[0], 4.0)

# Create inner table with 13 cells for JMB digits
inner_text = '|   ' * 13 + '|'
p_inner = jmb_label_table.rows[0].cells[1].paragraphs[0]
p_inner.clear()

# Let's create a simpler approach: a table with label column and 13 individual character cells
# Delete the previous tables and redo
# Actually, let's just use bordered cells properly

# Remove the incorrectly added tables
doc._body._body.remove(jmb_table._tbl)
doc._body._body.remove(jmb_label_table._tbl)

# 1.2 Matični broj/JMB
jmb_full = doc.add_table(rows=1, cols=15)
cell_text(jmb_full.rows[0].cells[0], '1.2 Matični broj/JMB', size=9, bold=True)
set_cell_width(jmb_full.rows[0].cells[0], 4.0)
set_cell_border(jmb_full.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
for j in range(1, 15):
    cell_text(jmb_full.rows[0].cells[j], '', size=9)
    set_cell_width(jmb_full.rows[0].cells[j], 0.75)
    set_cell_border(jmb_full.rows[0].cells[j],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
set_row_height(jmb_full.rows[0], 0.6)

# 1.3 Puni naziv / Prezime
name_table = doc.add_table(rows=1, cols=2)
cell_text(name_table.rows[0].cells[0], '1.3 Puni naziv / Prezime', size=9, bold=True)
set_cell_width(name_table.rows[0].cells[0], 4.5)
cell_text(name_table.rows[0].cells[1], '', size=9)
set_cell_width(name_table.rows[0].cells[1], 12.5)
set_cell_border(name_table.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
set_cell_border(name_table.rows[0].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})

# 1.4 Skraćeni naziv /Ime
short_name_table = doc.add_table(rows=1, cols=2)
cell_text(short_name_table.rows[0].cells[0], '1.4 Skraćeni naziv /Ime', size=9, bold=True)
set_cell_width(short_name_table.rows[0].cells[0], 4.5)
cell_text(short_name_table.rows[0].cells[1], '', size=9)
set_cell_width(short_name_table.rows[0].cells[1], 12.5)
set_cell_border(short_name_table.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
set_cell_border(short_name_table.rows[0].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})

# ---- Section 2: Adresa sjedišta / prebivališta ----
p = doc.add_paragraph()
run = p.add_run('2. Adresa sjedišta / prebivališta, odnosno boravišta')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)

# Address table
addr_labels = ['2.1 Država', '2.2 Opština', '2.3 Mjesto', '2.4 Ulica i broj',
               '2.5 Broj telefona/fax', '2.6 e-mail']
addr_table = doc.add_table(rows=len(addr_labels), cols=2)
for i, label in enumerate(addr_labels):
    cell_text(addr_table.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(addr_table.rows[i].cells[0], 4.0)
    cell_text(addr_table.rows[i].cells[1], '', size=9)
    set_cell_width(addr_table.rows[i].cells[1], 13.0)
    set_cell_border(addr_table.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(addr_table.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# ---- Section 3: Dodatak uz prijavu ----
p = doc.add_paragraph()
run = p.add_run('3. Dodatak uz prijavu:')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)

# 3.1
p = doc.add_paragraph()
run = p.add_run('     ☐  ')
run.font.size = Pt(11)
run = p.add_run('3.1 Dodatak A: Registracija pravnog lica')
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(4)

# 3.2
p = doc.add_paragraph()
run = p.add_run('     ☐  ')
run.font.size = Pt(11)
run = p.add_run('3.2 Dodatak B: Registracija fizičkog lica')
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(1)

# 3.2.1
p = doc.add_paragraph()
run = p.add_run('           3.2.1 Broj dodataka B     ')
run.font.size = Pt(9)
run2 = p.add_run('☐')
run2.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

# 3.3
p = doc.add_paragraph()
run = p.add_run('     ☐  ')
run.font.size = Pt(11)
run = p.add_run('3.3 Dodatak C: Evidencija ovlašćenih lica')
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(1)

# 3.3.1
p = doc.add_paragraph()
run = p.add_run('           3.3.1 Broj dodataka C     ')
run.font.size = Pt(9)
run2 = p.add_run('☐')
run2.font.size = Pt(12)
p.paragraph_format.space_after = Pt(6)

# ---- Section 4: Napomena ----
p = doc.add_paragraph()
run = p.add_run('4. Napomena')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_after = Pt(4)

# Napomena box
napomena_table = doc.add_table(rows=1, cols=1)
cell_text(napomena_table.rows[0].cells[0], '\n\n\n', size=10)
set_cell_border(napomena_table.rows[0].cells[0],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})

# ============================================================
# PAGE 2: Declaration + POPUNJAVA PORESKI ORGAN + Instructions
# ============================================================
add_page_break(doc)

# Declaration box
decl_table = doc.add_table(rows=3, cols=2)
# Row 0: Pod krivičnom...
a = decl_table.rows[0].cells[0]
b = decl_table.rows[0].cells[1]
a.merge(b)
cell_text(decl_table.rows[0].cells[0], 'Pod krivičnom odgovornošću izjavljujem da su podaci navedeni u prijavi tačni i potpuni.', size=9)
set_cell_border(decl_table.rows[0].cells[0],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# Row 1: JMB
jmb_row = decl_table.rows[1]
a = jmb_row.cells[0]
b = jmb_row.cells[1]
a.merge(b)
p = jmb_row.cells[0].paragraphs[0]
p.clear()
run = p.add_run('JMB:  ')
run.font.size = Pt(9)
# Add individual JMB boxes as text representation
run2 = p.add_run('|     ' * 13 + '|')
run2.font.size = Pt(9)
set_cell_border(jmb_row.cells[0],
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'},
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# Row 2: Signature and Date
sig_row = decl_table.rows[2]
cell_text(sig_row.cells[0], 'Potpis podnosioca / ovlašćenog lica:\n\n_______________________________', size=9)
set_cell_width(sig_row.cells[0], 10.0)
cell_text(sig_row.cells[1], 'Datum: _____/ _____/ _________', size=9)
set_cell_width(sig_row.cells[1], 7.0)
set_cell_border(sig_row.cells[0],
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
set_cell_border(sig_row.cells[1],
                right={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

doc.add_paragraph()  # spacing

# POPUNJAVA PORESKI ORGAN
p = doc.add_paragraph()
run = p.add_run('POPUNJAVA PORESKI ORGAN:')
run.font.size = Pt(10)
run.bold = True
p.paragraph_format.space_after = Pt(4)

po_table = doc.add_table(rows=3, cols=4)
po_table.style = 'Table Grid'
# Header row
headers = ['Broj dokumenta', 'Datum prijema', 'Datum obrade', 'Prezime i ime ovlašćenog službenika:']
for i, h in enumerate(headers):
    cell_text(po_table.rows[0].cells[i], h, size=8, bold=True)

# Empty rows for data
for r in range(1, 3):
    for c in range(4):
        cell_text(po_table.rows[r].cells[c], '', size=9)
    set_row_height(po_table.rows[r], 0.5)

doc.add_paragraph()  # spacing

# Instructions
p = doc.add_paragraph()
run = p.add_run('Uputstvo za popunjavanje obrasca JPR')
run.font.size = Pt(11)
run.bold = True
run.underline = True
p.paragraph_format.space_after = Pt(6)

instructions_1 = [
    ('1. Podaci o podnosiocu - ', 'U ovu rubriku se upisuju podaci o licu (pravnom ili fizičkom) koje podnosi prijavu.\nU polje 1.1 označava se tip lica - podnosioca sa znakom «X». Moguća je jedna oznaka, pravno lice (PL) ili fizičko lice (FL).\nU polje 1.2 unosi se matični broj pod kojim se PL vodi u Registru jedinica razvrstavanja (MONSTAT), odnosno za FL jedinstveni matični broj (JMB).\nU polje 1.3 unosi se puni naziv za PL, odnosno prezime za FL.\nU polje 1.4 unosi se skraćeni naziv za PL, odnosno ime za FL.'),
    ('2. Adresa sjedišta / prebivališta, odnosno boravišta', ' – U ovu rubriku se upisuju podaci o adresi sjedišta za PL, odnosno prebivališta za domaće fizičko lice ili podaci o adresi boravišta za strano fizičko lice (država, opština, mjesto, ulica i broj, broj telefona / fax i e-mail).'),
    ('3. Dodatak uz prijavu', ''),
]

for bold_text, normal_text in instructions_1:
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.font.size = Pt(9)
    run.bold = True
    if normal_text:
        run2 = p.add_run(normal_text)
        run2.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)

# Detailed instructions for section 3
instr_3_texts = [
    'Polje 3.1 se označava znakom «X» ukoliko se radi o registraciji pravnog lica koji popunjava i Dodatak «A» obrasca JPR.',
    'Polje 3.2 se označava znakom «X» ukoliko se radi o registraciji fizičkog lica (poreskog obveznika, odnosno osiguranika za penzijsko i invalidsko osiguranje) za koje se popunjava Dodatak «B» obrasca JPR. U polje 3.2.1 upisnje se broj dodataka B, odnosno za koliko lica je popunjen Dodatak «B».',
    'Polje 3.3 se označava znakom «X» ukoliko poreski obveznik (pravno lice, odnosno fizičko lice) prijavljuje ovlašćena lica (poreske punomoćnike). U polje 3.3.1 upisuje se broj dodataka C, odnosno za koliko lica je popunjen Dodatak «C».',
]
for txt in instr_3_texts:
    p = doc.add_paragraph()
    run = p.add_run(txt)
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)

# Section 4 instruction
p = doc.add_paragraph()
run = p.add_run('4. Napomena')
run.font.size = Pt(9)
run.bold = True
run2 = p.add_run(' – U ovu rubriku se upisuju prilozi, odnosno dokumenta koji se prilažu uz obrazac JPR, kao i ostali podaci od značaja za registraciju.')
run2.font.size = Pt(9)
p.paragraph_format.space_after = Pt(6)

# Additional instructions
additional_instr = [
    'Podnosilac, odnosno ovlašćeno lice za podnošenje prijave, tačnost i potpunost podataka potvrđuje izjavom. Takođe se unosi datum podnošenja prijave.',
    'Izmjena podataka vrši se podnošenjem obrasca JPR i odgovarajućeg dodatka, vezanog za nastalu promjenu i označava se znakom «X» u rubrici Izmjena podataka.',
    'Prestanak registracije (Brisanje iz registra) poreskog obveznika vrši se podnošenjem obrasca JPR i odgovarajućeg dodatka i označava se znakom «X» u rubrici Prestanak registracije (Brisanje iz registra).',
    'Prijava/odjava osiguranika vrši se podnošenjem obrasca JPR i Dodatak «B» i označava se znakom «X» u rubrici prijava/odjava osiguranika.',
    'Evidencija ovlašćenih lica vrši se podnošenjem obrasca JPR i Dodatak «C» i označava se znakom «X» u rubrici evidencija ovlašćenih lica.',
]
for txt in additional_instr:
    p = doc.add_paragraph()
    run = p.add_run(txt)
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)


# ============================================================
# PAGE 3: DODATAK A: Registracija pravnog lica
# ============================================================
add_page_break(doc)

p = doc.add_paragraph()
run = p.add_run('DODATAK A: Registracija pravnog lica')
run.font.size = Pt(14)
run.bold = True
p.paragraph_format.space_after = Pt(10)

# 1. Registracija
p = doc.add_paragraph()
run = p.add_run('1.    Registracija')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_after = Pt(4)

# 1.1 Naziv organa
reg_table = doc.add_table(rows=3, cols=2)
cell_text(reg_table.rows[0].cells[0], '1.1 Naziv organa', size=9, bold=True)
set_cell_width(reg_table.rows[0].cells[0], 4.5)
cell_text(reg_table.rows[0].cells[1], '', size=9)
set_cell_width(reg_table.rows[0].cells[1], 12.5)
set_cell_border(reg_table.rows[0].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})
set_cell_border(reg_table.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 1.2 Datum registracije - with individual cells
cell_text(reg_table.rows[1].cells[0], '1.2 Datum registracije', size=9, bold=True)
set_cell_width(reg_table.rows[1].cells[0], 4.5)
cell_text(reg_table.rows[1].cells[1], '____  /  ____  /  ________', size=9)
set_cell_width(reg_table.rows[1].cells[1], 12.5)
set_cell_border(reg_table.rows[1].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})
set_cell_border(reg_table.rows[1].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 1.3 Broj registarskog uloška
cell_text(reg_table.rows[2].cells[0], '1.3 Broj registarskog uloška', size=9, bold=True)
set_cell_width(reg_table.rows[2].cells[0], 4.5)
cell_text(reg_table.rows[2].cells[1], '', size=9)
set_cell_width(reg_table.rows[2].cells[1], 6.0)
set_cell_border(reg_table.rows[2].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})
set_cell_border(reg_table.rows[2].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 2. Drugi podaci
p = doc.add_paragraph()
run = p.add_run('2. Drugi podaci')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)

drugi_table = doc.add_table(rows=3, cols=2)
drugi_labels = ['2.1 Oblik organizovanja', '2.2 Oblik svojine', '2.3 Šifra djelatnosti']
for i, label in enumerate(drugi_labels):
    cell_text(drugi_table.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(drugi_table.rows[i].cells[0], 4.5)
    cell_text(drugi_table.rows[i].cells[1], '', size=9)
    set_cell_width(drugi_table.rows[i].cells[1], 12.5)
    set_cell_border(drugi_table.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(drugi_table.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 3. Podaci o osnivaču / vlasniku
p = doc.add_paragraph()
run = p.add_run('3. Podaci o osnivaču / vlasniku')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)

# Owner table
owner_table = doc.add_table(rows=4, cols=5)
owner_table.style = 'Table Grid'
# Headers
owner_headers = ['Jedinstveni matični broj / PIB', 'Prezime i ime / naziv', 'Adresa', '% vlasništva', 'X']
for i, h in enumerate(owner_headers):
    cell_text(owner_table.rows[0].cells[i], h, size=8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Set column widths
owner_widths = [4.5, 4.0, 4.0, 2.5, 1.0]
for row in owner_table.rows:
    for i, w in enumerate(owner_widths):
        set_cell_width(row.cells[i], w)

# Empty rows
for r in range(1, 4):
    for c in range(5):
        cell_text(owner_table.rows[r].cells[c], '', size=9)
    set_row_height(owner_table.rows[r], 0.7)

# 4. Podaci o odgovornom licu
p = doc.add_paragraph()
run = p.add_run('4. Podaci o odgovornom licu')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)

# 4.1 JMB
jmb_odg = doc.add_table(rows=1, cols=15)
cell_text(jmb_odg.rows[0].cells[0], '4.1 Jedinstveni matični broj', size=9, bold=True)
set_cell_width(jmb_odg.rows[0].cells[0], 5.0)
set_cell_border(jmb_odg.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
for j in range(1, 15):
    cell_text(jmb_odg.rows[0].cells[j], '', size=9)
    set_cell_width(jmb_odg.rows[0].cells[j], 0.75)
    set_cell_border(jmb_odg.rows[0].cells[j],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})

# 4.2 Prezime, 4.3 Ime, 4.4 Adresa
odg_fields = doc.add_table(rows=3, cols=2)
odg_labels = ['4.2 Prezime', '4.3 Ime', '4.4 Adresa']
for i, label in enumerate(odg_labels):
    cell_text(odg_fields.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(odg_fields.rows[i].cells[0], 4.0)
    cell_text(odg_fields.rows[i].cells[1], '', size=9)
    set_cell_width(odg_fields.rows[i].cells[1], 13.0)
    set_cell_border(odg_fields.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(odg_fields.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# Instructions for Dodatak A
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Uputstvo za popunjavanje Dodatka „A": Registracija pravnog lica')
run.font.size = Pt(10)
run.bold = True
run.underline = True
p.paragraph_format.space_after = Pt(6)

dodatak_a_instr = [
    '1. Registracija – U ovu rubriku se upisuju podaci o registraciji pravnog lica (naziv organa kod kojeg je registrovano pravno lice, datum registracije i broj registarskog uloška).',
    '2. Drugi podaci – U ovu rubriku se upisuju podaci o pravnom licu (šifra oblika organizovanja, šifra oblika svojine i šifra djelatnosti) koji se nalaze u Obavještenju o razvrstavanju po djelatnostima (MONSTAT).',
    '3. Podaci o osnivaču/vlasniku – U ovu rubriku se unose podaci o osnivaču/vlasniku pravnog lica (matični broj fizičkog lica / PIB pravnog lica; prezime i ime / naziv; adresa i procenat vlasništva). Ukoliko je došlo do prestanka vlasništva određenog lica potrebno je označiti znakom «X» navedenu promjenu.',
    '4. Podaci o odgovornom licu – U ovu rubriku se upisuju podaci o odgovornom licu, odnosno licu ovlašćenom za zastupanje (matični broj, prezime i ime i adresa).',
]
for txt in dodatak_a_instr:
    p = doc.add_paragraph()
    run = p.add_run(txt)
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)


# ============================================================
# PAGE 4: DODATAK B: Registracija fizičkog lica
# ============================================================
add_page_break(doc)

p = doc.add_paragraph()
run = p.add_run('DODATAK B: Registracija fizičkog lica')
run.font.size = Pt(14)
run.bold = True
p.paragraph_format.space_after = Pt(8)

# 1. Opšti podaci
p = doc.add_paragraph()
run = p.add_run('1. Opšti podaci')
run.font.size = Pt(10)
run.bold = True
p.paragraph_format.space_after = Pt(4)

# 1.1 JMB
jmb_b = doc.add_table(rows=1, cols=15)
cell_text(jmb_b.rows[0].cells[0], '1.1 JMB', size=9, bold=True)
set_cell_width(jmb_b.rows[0].cells[0], 3.5)
set_cell_border(jmb_b.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
for j in range(1, 15):
    cell_text(jmb_b.rows[0].cells[j], '', size=9)
    set_cell_width(jmb_b.rows[0].cells[j], 0.75)
    set_cell_border(jmb_b.rows[0].cells[j],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})

# 1.2 Prezime + 1.3 Ime on same row
name_b = doc.add_table(rows=1, cols=4)
cell_text(name_b.rows[0].cells[0], '1.2  Prezime', size=9, bold=True)
set_cell_width(name_b.rows[0].cells[0], 3.5)
cell_text(name_b.rows[0].cells[1], '', size=9)
set_cell_width(name_b.rows[0].cells[1], 5.5)
cell_text(name_b.rows[0].cells[2], '1.3 Ime', size=9, bold=True)
set_cell_width(name_b.rows[0].cells[2], 2.0)
cell_text(name_b.rows[0].cells[3], '', size=9)
set_cell_width(name_b.rows[0].cells[3], 5.5)
for c in [0, 2]:
    set_cell_border(name_b.rows[0].cells[c],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
for c in [1, 3]:
    set_cell_border(name_b.rows[0].cells[c],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})

# 1.4 Ime jednog roditelja + 1.5 Djevojačko prezime
parent_b = doc.add_table(rows=1, cols=4)
cell_text(parent_b.rows[0].cells[0], '1.4  Ime jednog roditelja', size=9, bold=True)
set_cell_width(parent_b.rows[0].cells[0], 4.5)
cell_text(parent_b.rows[0].cells[1], '', size=9)
set_cell_width(parent_b.rows[0].cells[1], 4.0)
cell_text(parent_b.rows[0].cells[2], '1.5 Djevojačko prezime', size=9, bold=True)
set_cell_width(parent_b.rows[0].cells[2], 4.0)
cell_text(parent_b.rows[0].cells[3], '', size=9)
set_cell_width(parent_b.rows[0].cells[3], 4.0)
for c in [0, 2]:
    set_cell_border(parent_b.rows[0].cells[c],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
for c in [1, 3]:
    set_cell_border(parent_b.rows[0].cells[c],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})

# 1.6 Datum rođenja
dob_table = doc.add_table(rows=1, cols=2)
cell_text(dob_table.rows[0].cells[0], '1.6 Datum rođenja', size=9, bold=True)
set_cell_width(dob_table.rows[0].cells[0], 3.5)
cell_text(dob_table.rows[0].cells[1], '____  .  ____  .  ________', size=9)
set_cell_width(dob_table.rows[0].cells[1], 6.0)
set_cell_border(dob_table.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
set_cell_border(dob_table.rows[0].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})

# Fields 1.7 - 1.9 
fields_1_7to9 = doc.add_table(rows=3, cols=2)
labels_7to9 = ['1.7 Mjesto rođenja', '1.8 Opština rođenja', '1.9 Država rođenja']
for i, label in enumerate(labels_7to9):
    cell_text(fields_1_7to9.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(fields_1_7to9.rows[i].cells[0], 3.5)
    cell_text(fields_1_7to9.rows[i].cells[1], '', size=9)
    set_cell_width(fields_1_7to9.rows[i].cells[1], 13.5)
    set_cell_border(fields_1_7to9.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(fields_1_7to9.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 1.10 Pol
pol_table = doc.add_table(rows=1, cols=5)
cell_text(pol_table.rows[0].cells[0], '1.10 Pol', size=9, bold=True)
set_cell_width(pol_table.rows[0].cells[0], 3.5)
cell_text(pol_table.rows[0].cells[1], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_width(pol_table.rows[0].cells[1], 0.8)
cell_text(pol_table.rows[0].cells[2], 'Muški', size=9, bold=True)
set_cell_width(pol_table.rows[0].cells[2], 4.0)
cell_text(pol_table.rows[0].cells[3], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_width(pol_table.rows[0].cells[3], 0.8)
cell_text(pol_table.rows[0].cells[4], 'Ženski', size=9, bold=True)
set_cell_width(pol_table.rows[0].cells[4], 4.0)
for cell in pol_table.rows[0].cells:
    set_cell_border(cell,
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 1.11 Državljanstvo
drz_table = doc.add_table(rows=1, cols=2)
cell_text(drz_table.rows[0].cells[0], '1.11 Državljanstvo', size=9, bold=True)
set_cell_width(drz_table.rows[0].cells[0], 3.5)
cell_text(drz_table.rows[0].cells[1], '', size=9)
set_cell_width(drz_table.rows[0].cells[1], 13.5)
set_cell_border(drz_table.rows[0].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})
set_cell_border(drz_table.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 1.12 Vrsta identifikacionog dokumenta + 1.13 Broj identifikacionog dokumenta
id_table = doc.add_table(rows=1, cols=4)
cell_text(id_table.rows[0].cells[0], '1.12 Vrsta identifikacionog\n       dokumenta', size=9, bold=True)
set_cell_width(id_table.rows[0].cells[0], 4.5)
cell_text(id_table.rows[0].cells[1], '', size=9)
set_cell_width(id_table.rows[0].cells[1], 3.5)
cell_text(id_table.rows[0].cells[2], '1.13 Broj identifikacionog\n       dokumenta', size=9, bold=True)
set_cell_width(id_table.rows[0].cells[2], 4.5)
cell_text(id_table.rows[0].cells[3], '', size=9)
set_cell_width(id_table.rows[0].cells[3], 4.0)
for c in [0, 2]:
    set_cell_border(id_table.rows[0].cells[c],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
for c in [1, 3]:
    set_cell_border(id_table.rows[0].cells[c],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})

# 1.14 - 1.16
fields_14to16 = doc.add_table(rows=3, cols=2)
labels_14to16 = ['1.14 Izdat od', '1.15 Nivo obrazovanja', '1.16 Kvalifikacija']
for i, label in enumerate(labels_14to16):
    cell_text(fields_14to16.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(fields_14to16.rows[i].cells[0], 3.5)
    cell_text(fields_14to16.rows[i].cells[1], '', size=9)
    set_cell_width(fields_14to16.rows[i].cells[1], 13.5)
    set_cell_border(fields_14to16.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(fields_14to16.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 2. Adresa prebivališta, odnosno boravišta
p = doc.add_paragraph()
run = p.add_run('2.    Adresa prebivališta, odnosno boravišta')
run.font.size = Pt(10)
run.bold = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)

addr_b_labels = ['2.1 Država', '2.2 Opština', '2.3 Mjesto', '2.4 Ulica i broj',
                 '2.5 Broj telefona / fax/ e-mail']
addr_b_table = doc.add_table(rows=len(addr_b_labels), cols=2)
for i, label in enumerate(addr_b_labels):
    cell_text(addr_b_table.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(addr_b_table.rows[i].cells[0], 4.5)
    cell_text(addr_b_table.rows[i].cells[1], '', size=9)
    set_cell_width(addr_b_table.rows[i].cells[1], 12.5)
    set_cell_border(addr_b_table.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(addr_b_table.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 3. Registracija poreskog obveznika
p = doc.add_paragraph()
run = p.add_run('3.  Registracija poreskog obveznika')
run.font.size = Pt(10)
run.bold = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)

# 3.1 Samostalna djelatnost
sam_djel = doc.add_table(rows=1, cols=6)
cell_text(sam_djel.rows[0].cells[0], '3.1  Samostalna djelatnost', size=9, bold=True)
set_cell_width(sam_djel.rows[0].cells[0], 4.5)
cell_text(sam_djel.rows[0].cells[1], '3.1.1  Stalna', size=9, bold=True)
set_cell_width(sam_djel.rows[0].cells[1], 3.0)
cell_text(sam_djel.rows[0].cells[2], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_width(sam_djel.rows[0].cells[2], 0.8)
cell_text(sam_djel.rows[0].cells[3], '3.1.2  Vremenski ograničena', size=9, bold=True)
set_cell_width(sam_djel.rows[0].cells[3], 4.5)
cell_text(sam_djel.rows[0].cells[4], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_width(sam_djel.rows[0].cells[4], 0.8)
cell_text(sam_djel.rows[0].cells[5], '', size=9)
set_cell_width(sam_djel.rows[0].cells[5], 2.0)
for cell in sam_djel.rows[0].cells:
    set_cell_border(cell,
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 3.1.3 Registracija table
p = doc.add_paragraph()
run = p.add_run('3.1.3 Registracija')
run.font.size = Pt(9)
run.bold = True
p.paragraph_format.space_after = Pt(2)

reg_b_table = doc.add_table(rows=2, cols=3)
reg_b_table.style = 'Table Grid'
cell_text(reg_b_table.rows[0].cells[0], 'Naziv organa', size=8, bold=True)
cell_text(reg_b_table.rows[0].cells[1], 'Datum registracije', size=8, bold=True)
cell_text(reg_b_table.rows[0].cells[2], 'Broj registarskog uloška', size=8, bold=True)
for c in range(3):
    cell_text(reg_b_table.rows[1].cells[c], '', size=9)
set_row_height(reg_b_table.rows[1], 0.5)

# 3.1.4 Djelatnost
djel_table = doc.add_table(rows=1, cols=4)
cell_text(djel_table.rows[0].cells[0], '3.1.4 Djelatnost', size=9, bold=True)
set_cell_width(djel_table.rows[0].cells[0], 3.0)
cell_text(djel_table.rows[0].cells[1], 'Opis', size=9, bold=True)
set_cell_width(djel_table.rows[0].cells[1], 6.0)
cell_text(djel_table.rows[0].cells[2], '', size=9)
set_cell_width(djel_table.rows[0].cells[2], 4.0)
cell_text(djel_table.rows[0].cells[3], 'Šifra', size=9, bold=True)
set_cell_width(djel_table.rows[0].cells[3], 3.0)
for c in [0, 1, 3]:
    set_cell_border(djel_table.rows[0].cells[c],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
set_cell_border(djel_table.rows[0].cells[2],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})

# 3.1.5 Mjesto poslovanja
p = doc.add_paragraph()
run = p.add_run('3.1.5 Mjesto poslovanja')
run.font.size = Pt(9)
run.bold = True
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.space_before = Pt(4)

mj_posl_labels = ['3.1.5.1 Naziv', '3.1.5.2 Država', '3.1.5.3 Opština', '3.1.5.4 Mjesto',
                   '3.1.5.5 Ulica i broj', '3.1.5.6 Broj telefona']
mj_posl = doc.add_table(rows=len(mj_posl_labels), cols=2)
for i, label in enumerate(mj_posl_labels):
    cell_text(mj_posl.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(mj_posl.rows[i].cells[0], 4.0)
    cell_text(mj_posl.rows[i].cells[1], '', size=9)
    set_cell_width(mj_posl.rows[i].cells[1], 13.0)
    set_cell_border(mj_posl.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(mj_posl.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 4. Registracija osiguranika za PIO
p = doc.add_paragraph()
run = p.add_run('4.  Registracija osiguranika za PIO')
run.font.size = Pt(10)
run.bold = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)

pio_labels = [
    '4.1 Opština zaposlenja/obavljanja\n      djelatnosti',
    '4.2 Osnov osiguranja',
    '4.3 Radno vrijeme',
    '4.4 Osnov prestanka osiguranja',
    '4.5 Datum'
]
pio_table = doc.add_table(rows=len(pio_labels), cols=2)
for i, label in enumerate(pio_labels):
    cell_text(pio_table.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(pio_table.rows[i].cells[0], 5.5)
    if label == '4.5 Datum':
        cell_text(pio_table.rows[i].cells[1], '____  .  ____  .  ________', size=9)
    else:
        cell_text(pio_table.rows[i].cells[1], '', size=9)
    set_cell_width(pio_table.rows[i].cells[1], 11.5)
    set_cell_border(pio_table.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(pio_table.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# ============================================================
# PAGE 5: Continuation of Dodatak B (4.6-4.8) + Instructions
# ============================================================
add_page_break(doc)

# 4.6, 4.7
fields_46_47 = doc.add_table(rows=2, cols=2)
labels_46_47 = ['4.6 Nivo obrazovanja za radno mjesto', '4.7 Kvalifikacija za radno mjesto']
for i, label in enumerate(labels_46_47):
    cell_text(fields_46_47.rows[i].cells[0], label, size=9, bold=True)
    set_cell_width(fields_46_47.rows[i].cells[0], 6.5)
    cell_text(fields_46_47.rows[i].cells[1], '', size=9)
    set_cell_width(fields_46_47.rows[i].cells[1], 10.5)
    set_cell_border(fields_46_47.rows[i].cells[1],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})
    set_cell_border(fields_46_47.rows[i].cells[0],
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 4.8 Vrsta zaposlenja
vz_table = doc.add_table(rows=1, cols=5)
cell_text(vz_table.rows[0].cells[0], '4.8 Vrsta zaposlenja:', size=9, bold=True)
set_cell_width(vz_table.rows[0].cells[0], 4.0)
cell_text(vz_table.rows[0].cells[1], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_width(vz_table.rows[0].cells[1], 0.8)
cell_text(vz_table.rows[0].cells[2], '4.8.1 Neodređeno', size=9)
set_cell_width(vz_table.rows[0].cells[2], 4.0)
cell_text(vz_table.rows[0].cells[3], '☐', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_width(vz_table.rows[0].cells[3], 0.8)
cell_text(vz_table.rows[0].cells[4], '4.8.2 Određeno', size=9)
set_cell_width(vz_table.rows[0].cells[4], 4.0)
for cell in vz_table.rows[0].cells:
    set_cell_border(cell,
                    top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# Instructions for Dodatak B
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Uputstvo za popunjavanje Dodatka „B": Registracija fizičkog lica')
run.font.size = Pt(10)
run.bold = True
run.underline = True
p.paragraph_format.space_after = Pt(6)

b_instructions = [
    ('1. Opšti podaci', ' – U ovu rubriku se upisuju podaci o poreskom obvezniku, odnosno osiguraniku (jedinstveni matični broj, prezime, ime, ime jednog roditelja, djevojačko prezime, datum rođenja, mjesto rođenja, opština rođenja, država rođenja, pol, državljanstvo). Takođe se upisuju vrsta i broj identifikacionog dokumenta i naziv organa koji ga je izdao (država se upisuje ukoliko je identifikacioni dokument pasoš).'),
    ('', 'U polje 1.15 upisuje se nivo obrazovanja, odnosno promjeni podatka o nivou obrazovanja.\nU polje 1.16 upisuje se kvalifikacija, odnosno promjeni podatka o kvalifikaciji.'),
    ('2. Adresa prebivališta, odnosno boravišta', ' – U ovu rubriku upisuju se podaci o adresi prebivališta za domaće fizičko lice ili podaci o adresi boravišta za strano fizičko lice (država, opština, mjesto, ulica i broj, broj telefona / fax i e-mail).'),
    ('3. Registracija poreskog obveznika', ' – U ovu rubriku se označava polje «X» ako je poreski obveznik fizičko lice koje obavlja samostalnu djelatnost.'),
    ('3.1.Samostalna djelatnost – Polje 3.1.1', ' se označava znakom «X» ako se radi o fizičkom licu koje neprekidno obavlja djelatnost po osnovu koje je i registrovan, odnosno polje 3.1.2 se označava znakom «X» ako se radi o fizičkom licu  koji vremenski ograničeno obavlja djelatnost na osnovu rješenja nadležnog organa.'),
    ('', 'U polje 3.1.3. upisuju se podaci o registraciji lica koje obavlja samostalnu djelatnost (naziv organa kod kojeg je registrovan, datum registracije, broj registarskog uloška).\nU polje 3.1.4. upisuje se opis djelatnosti, odnosno šifra djelatnosti prema klasifikaciji djelatnosti.\nU polje 3.1.5. upisuje se podaci o Mjestu poslovanja - objektu (kancelariji) u kojoj poreski obveznik obavlja samostalnu djelatnost (nazivi; država; opština; mjesto, ulica i broj i broj telefona).'),
    ('4. Registracija osiguranika za penzijsko i invalidsko osiguranje (PIO)', ' – U ovu rubriku se označava polje «X» ako se fizičko lice osigurava za penzijsko i invalidsko osiguranje.'),
    ('4.1 Opština zaposlenja/obavljanja djelatnosti', ' – U ovu rubriku upisuje se naziv opštine u kojoj je osigurnaik zaposlen, odnosno obavlja djelatnosti.'),
    ('4.2 Osnov osiguranja', ' – U ovu rubriku upisuje se opisno osnov osiguranja i šifra iz Šifrarnika osnova za penzijsko i invalidsko osiguranje, koji je odštampan uz ovo uputstvo i čini njegov sastavni dio.'),
    ('4.3 Radno vrijeme osiguranika', ' – U ovu rubriku upisuje se nedeljno radno vrijeme (40 časova) ili onoliko časova na koliko se osiguranik prijavljuje, ali koje ne može biti manje od 10 časova nedeljno, izuzev za rad direktora.'),
    ('4.4 Osnov prestanka osiguranja', ' – U ovu rubriku upisuje se šifra i naziv osnova prestanka osiguranja i to: 1 – prestanak zaposlenja; 5 – smrt; 6 – ostvarivanje prava na penziju; 7 – drugi razlozi; 8 – prestanak obavljanja samostalne djelatnosti; 9 – prestanak rada agenvernih poslova; 10 – prestanak obavljanja poljoprivredne djelatnosti; 11 – prestanak ostvarivanja novčane naknade koja osigurava nezaposlena lica od Zavoda za zapošljavanje.'),
    ('4.5 Datum', ' – u ovu rubriku se upisuje datum sticanja svojstva osiguranika ili datum prestanka sticanja svojstva osiguranika u zavisnosti od podnošenja podataka za osiguranika PIO-a u prijavi. Podatak se upisuje iz ugovora o radu, akta nadležnog organa, odnosno rješenja kojim je utvrđen osnov prestanka osiguranika i drugih akata u skladu sa zakonom.'),
    ('', 'U polje 4.6 upisuje se nivo obrazovanja za radno mjesto na koje je lice zaposleno.\nU polje 4.7 upisuje se kvalifikacija za radno mjesto na koje je lice zaposleno.'),
    ('4.8 Vrsta zaposlenja', ' – Polje 4.8.1 se označava znakom «X» ako se radi o zaposlenom licu na neodređeno vrijeme, odnosno polje 4.8.2 se označava znakom «X» ako se radi o zaposlenom licu na određeno vrijeme.'),
]
for bold_text, normal_text in b_instructions:
    p = doc.add_paragraph()
    if bold_text:
        run = p.add_run(bold_text)
        run.font.size = Pt(8)
        run.bold = True
    if normal_text:
        run2 = p.add_run(normal_text)
        run2.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(3)

# ============================================================
# PAGE 6: ŠIFRARNIK OSNOVA ZA PENZIJSKO I INVALIDSKO OSIGURANJE
# ============================================================
add_page_break(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ŠIFRARNIK OSNOVA ZA  PENZIJSKO I INVALIDSKO OSIGURANJE')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_after = Pt(8)

# The big table
sifrarnik_data = [
    ('OSNOV  OSIGURANJA', 'Šifra'),
    ('ZA OSIGURANIKE ZAPOSLENE', ''),
    ('lica zaposlena u privrednom društvu, drugom pravnom licu, državnom organu, organu jedinice lokalne samouprave ili kod fizičkog lica (poslodavac)', '101'),
    ('civilna lica na službi u Vojsci Crne Gore', '136'),
    ('profesionalna vojna lica na službi u Vojsci Crne Gore', '408'),
    ('izabrana ili imenovana lica, ako za obavljanje funkcije ostvaruju zaradu', '102'),
    ('članovi upravnih odbora u javnim preduzećima i ustanovama koji za svoj rad primaju naknadu ako nijesu osigurani po drugom osnovu', '157'),
    ('članovi odbora direktora u privrednom društvu i drugom pravnom licu koji za svoj rad primaju naknadu ako nijesu osigurani po drugom osnovu', '158'),
    ('zaposlena lica upućena na rad u inostranstvo, odnosno zaposlena u privrednom društvu, odnosno drugom pravnom licu koja obavlja djelatnost ili usluge u inostranstvu, ako nijesu obavezno osigurana po propisima te zemlje ili ako međunarodnim ugovorom nije drugačije određeno', '129'),
    ('državljani Crne Gore koji su na teritoriji Crne Gore zaposleni kod stranih ili međunarodnih organizacija i ustanova, stranih diplomatskih i konzularnih predstavništava ili kod stranih pravnih i fizičkih lica, ako međunarodnim ugovorom nije drugačije određeno', '130'),
    ('državljani Crne Gore zaposleni u inostranstvu, ako za to vrijeme nijesu obavezno osigurani kod stranog nosioca osiguranja ili ako prava iz penzijskog i invalidskog osiguranja, po propisima te države, ne mogu ostvariti ili koristiti van njene teritorije, a neposredno prije odlaska u inostranstvo bili su osigurani u Crnoj Gori, odnosno ako su imali prije odlaska u inostranstvo stalno prebivalište u Crnoj Gori', '103'),
    ('strani državljani i lica bez državljanstva koji su na teritoriji Crne Gore zaposleni kod stranih pravnih i fizičkih lica, ako međunarodnim ugovorom nije drugačije određeno, odnosno ako nijesu osigurani po propisima druge države', '137'),
    ('strani državljani i lica državljanstva koji su na teritoriji Crne Gore zaposleni kod međunarodnih organizacija i ustanova, ako te osiguranike predviđeno međunarodnim ugovorom', '139'),
    ('lica koja odsustvuju sa rada do navršene treće godine života djeteta', '156'),
    ('lica koja obavljaju povremene i privremene poslove', '121'),
    ('lica koja obavljaju poslove van prostorija poslodavca', '138'),
    ('lica koja rade sa nepunim radnim vremenom, u skladu sa Zakonom o radu', '409'),
    ('lica koja rade sa skraćenim radnim vremenom, zbog otežanih uslova rada', '161'),
    ('korisnik djelimične invalidske penzije koji radi ¼ punog radnog vremena', '165'),
    ('korisnik starosne penzije koji je zaposlen', '160'),
    ('lica koja rade sa polovinom punog radnog vremena zbog pojačane njege djeteta do navršene treće godine života djeteta ili u njeze i staranja djeteta usljed teških oboljenja', '164'),
    ('korisnik porodične penzije (dijete) koji je zaposlen', '169'),
    ('korisnik vojne invalidske penzije koji je zaposlen', '170'),
    ('korisnik starosne penzije koji je član odbora direktora u privrednom društvu i drugom pravnom licu, odnosno član upravnog odbora u javnim preduzećima i ustanovama koji za svoj rad prima naknadu a nije osiguran po drugom osnovu', '170'),
    ('korisnik djelimične invalidske penzije koji je član odbora direktora u privrednom društvu i drugom pravnom licu, odnosno član upravnog odbora u javnim preduzećima i ustanovama koji za svoj rad prima naknadu a nije osiguran po drugom osnovu', '170'),
    ('ZA OSIGURANIKE KOJI OBAVLJAJU SAMOSTALNU DJELATNOST', ''),
    ('lica koja samostalno obavljaju privrednu ili drugu djelatnost, a nijesu obavezno osigurana po osnovu zaposlenja', '232'),
    ('lica koja obavljaju poslove po osnovu ugovora o djelu, a nijesu osigurani po drugom osnovu', '263'),
    ('lica koja obavljaju poslove po osnovu autorskog ugovora, a nijesu osigurani po drugom osnovu', '266'),
    ('lica koja obavljaju poslove po osnovu drugih ugovora kod kojih za izvršeni posao ostvaruju naknadu, a nijesu osigurani po drugom osnovu', '267'),
    ('sveštenici, vjerski službenici, monasi i monahinje, ako nijesu obavezno osigurani po drugom osnovu', '250'),
    ('korisnik starosne penzije koji obavlja i samostalnu djelatnost', '269'),
    ('korisnik porodične penzije (dijete) koji obavlja i samostalnu djelatnost', '270'),
    ('korisnik porodične penzije koji obavlja i poslove po osnovu ugovora o djelu, odnosno poslove po osnovu autorskog ugovora, kao i poslove po osnovu drugih ugovora, kod kojih za izvršen posao ostvaruje ugovorenu naknadu', '271'),
    ('korisnik djelimične invalidske penzije koji obavlja i samostalnu djelatnost', '272'),
    ('korisnik vojne invalidske penzije koji obavlja i samostalnu djelatnost', '273'),
    ('ZA OSIGURANIKE KOJI OBAVLJAJU POLJOPRIVREDNU DJELATNOST', ''),
    ('poljoprivrednici', '320'),
    ('LICA ZA ČIJIM JE RADOM PRESTALA POTREBA, ODNOSNO KOJA SU PRESTALA DA OBAVLJAJU PREDUZETNIČKU DJELATNOST DOK OSTVARUJU NOVČANU NAKNADU PREMA PROPISIMA O RADU I ZAPOŠLJAVANJU', '155'),
    ('PRODUŽENO OSIGURANJE', '407'),
    ('korisnice prava na naknadu po osnovu ranijeg korišćenog prava na naknadu po osnovu rođenja troje ili više djece kojima je radi korišćenja tog prava njihovom voljom prestao radni odnos na neodređeno vrijeme', '410'),
]

sif_table = doc.add_table(rows=len(sifrarnik_data), cols=2)
sif_table.style = 'Table Grid'

for i, (text, sifra) in enumerate(sifrarnik_data):
    cell_text(sif_table.rows[i].cells[0], text, size=7, bold=(i <= 1 or sifra == ''))
    cell_text(sif_table.rows[i].cells[1], sifra, size=7, bold=(i == 0), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_width(sif_table.rows[i].cells[0], 15.0)
    set_cell_width(sif_table.rows[i].cells[1], 2.0)


# ============================================================
# PAGE 7: DODATAK C: Evidencija ovlašćenih lica
# ============================================================
add_page_break(doc)

p = doc.add_paragraph()
run = p.add_run('DODATAK C: Evidencija ovlašćenih lica')
run.font.size = Pt(14)
run.bold = True
p.paragraph_format.space_after = Pt(10)

# 1. Tip ovlašćenja
p = doc.add_paragraph()
run = p.add_run('1. Tip ovlašćenja ')
run.font.size = Pt(11)
run.bold = True
run2 = p.add_run('(Upisati \'X\' pored jednog tipa ovlašćenja):')
run2.font.size = Pt(9)
p.paragraph_format.space_after = Pt(4)

# Checkboxes
for cb_text in ['1.1 Interno ovlašćenje', '1.2 Ugovorno ovlašćenje', '1.3 Ostalo ovlašćenje']:
    p = doc.add_paragraph()
    run = p.add_run('     ☐  ')
    run.font.size = Pt(11)
    run2 = p.add_run(cb_text)
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# 2. Podaci o ovlašćenom licu
p = doc.add_paragraph()
run = p.add_run('2.  Podaci o ovlašćenom licu')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)

# 2.1 JMB
jmb_c = doc.add_table(rows=1, cols=15)
cell_text(jmb_c.rows[0].cells[0], '2.1  JMB', size=9, bold=True)
set_cell_width(jmb_c.rows[0].cells[0], 3.5)
set_cell_border(jmb_c.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})
for j in range(1, 15):
    cell_text(jmb_c.rows[0].cells[j], '', size=9)
    set_cell_width(jmb_c.rows[0].cells[j], 0.75)
    set_cell_border(jmb_c.rows[0].cells[j],
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})

# 2.2 Prezime i ime
c_fields_1 = doc.add_table(rows=1, cols=2)
cell_text(c_fields_1.rows[0].cells[0], '2.2 Prezime i ime', size=9, bold=True)
set_cell_width(c_fields_1.rows[0].cells[0], 3.5)
cell_text(c_fields_1.rows[0].cells[1], '', size=9)
set_cell_width(c_fields_1.rows[0].cells[1], 13.5)
set_cell_border(c_fields_1.rows[0].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})
set_cell_border(c_fields_1.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 2.3 Poreski obveznik
p = doc.add_paragraph()
run = p.add_run('2.3 Poreski obveznik ')
run.font.size = Pt(9)
run.bold = True
run2 = p.add_run('(popunjava se ako je označeno 1.2. Ugovorno ovlašćenje)')
run2.font.size = Pt(8)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)

# 2.3.1 PIB
c_pib = doc.add_table(rows=1, cols=2)
cell_text(c_pib.rows[0].cells[0], '          2.3.1 PIB', size=9, bold=True)
set_cell_width(c_pib.rows[0].cells[0], 4.0)
cell_text(c_pib.rows[0].cells[1], '', size=9)
set_cell_width(c_pib.rows[0].cells[1], 6.0)
set_cell_border(c_pib.rows[0].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})
set_cell_border(c_pib.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 2.3.2 Naziv
c_naziv = doc.add_table(rows=1, cols=2)
cell_text(c_naziv.rows[0].cells[0], '          2.3.2 Naziv', size=9, bold=True)
set_cell_width(c_naziv.rows[0].cells[0], 4.0)
cell_text(c_naziv.rows[0].cells[1], '', size=9)
set_cell_width(c_naziv.rows[0].cells[1], 6.0)
set_cell_border(c_naziv.rows[0].cells[1],
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'})
set_cell_border(c_naziv.rows[0].cells[0],
                top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

# 3. Vrsta ovlašćenja
p = doc.add_paragraph()
run = p.add_run('3.  Vrsta ovlašćenja (moguće izabrati jednu ili više tipova ovlašćenja):')
run.font.size = Pt(11)
run.bold = True
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(6)

# 3.1 - 3.4 sections
ovl_sections = [
    '3.1 Podnošenje zahtjeva',
    '3.2 Podnošenje i pregled poreskih prijava',
    '3.3  Pregled analitičke kartice',
    '3.4  Pristup e-sandučetu',
]

for section_title in ovl_sections:
    p = doc.add_paragraph()
    run = p.add_run(section_title)
    run.font.size = Pt(10)
    run.bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

    # Prijava checkbox
    p = doc.add_paragraph()
    run = p.add_run('     ☐')
    run.font.size = Pt(11)
    run2 = p.add_run('Prijava')
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

    # Period trajanja ovlašćenja
    period_table = doc.add_table(rows=1, cols=5)
    cell_text(period_table.rows[0].cells[0], '          Period trajanja ovlašćenja:', size=8)
    set_cell_width(period_table.rows[0].cells[0], 5.0)
    cell_text(period_table.rows[0].cells[1], '☐', size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_width(period_table.rows[0].cells[1], 0.6)
    cell_text(period_table.rows[0].cells[2], 'Neograničeno', size=8)
    set_cell_width(period_table.rows[0].cells[2], 3.0)
    cell_text(period_table.rows[0].cells[3], '☐', size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_width(period_table.rows[0].cells[3], 0.6)
    cell_text(period_table.rows[0].cells[4], 'Vremenski ograničeno do dana', size=8)
    set_cell_width(period_table.rows[0].cells[4], 6.0)
    # Add a field for the date
    p_date = period_table.rows[0].cells[4].paragraphs[0]
    run_field = p_date.add_run('  _______________')
    run_field.font.size = Pt(8)
    for cell in period_table.rows[0].cells:
        set_cell_border(cell,
                        top={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                        bottom={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                        left={'val': 'none', 'sz': '0', 'color': 'FFFFFF'},
                        right={'val': 'none', 'sz': '0', 'color': 'FFFFFF'})

    # Prestanak checkbox
    p = doc.add_paragraph()
    run = p.add_run('     ☐')
    run.font.size = Pt(11)
    run2 = p.add_run(' Prestanak')
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


# ============================================================
# PAGE 8: Instructions for Dodatak C + Član 5
# ============================================================
add_page_break(doc)

p = doc.add_paragraph()
run = p.add_run('Uputstvo za popunjavanje Dodatka „C": Evidencija ovlašćenih lica')
run.font.size = Pt(10)
run.bold = True
run.underline = True
p.paragraph_format.space_after = Pt(6)

# Intro paragraph
p = doc.add_paragraph()
run = p.add_run('Dodatak C popunjava poreski obveznik (pravno, odnosno fizičko lice) koji pismeno ovlasti lice (poreskog punomoćnika) da u njegovo ime vodi poslove u vezi sa izvršavanjem poreskih obaveza u skladu sa poreskim propisima. Lice koje je ovlašćeno da zastupa poreskog obveznika kod poreskog organa mora biti kao takvo evidentirano u Poreskom registru.')
run.font.size = Pt(9)
p.paragraph_format.space_after = Pt(6)

# Instruction items
c_instr = [
    ('1.', 'Tip ovlašćenja', ' – U ovu rubriku se označava znakom «X» Tip ovlašćenja: 1.1 Interno ovlašćenje; 1.2 Ugovorno ovlašćenje i 1.3 Ostalo ovlašćenje.'),
]
for num, bold_text, normal_text in c_instr:
    p = doc.add_paragraph()
    run = p.add_run(f'{num} ')
    run.font.size = Pt(9)
    run.bold = True
    run2 = p.add_run(bold_text)
    run2.font.size = Pt(9)
    run2.bold = True
    run3 = p.add_run(normal_text)
    run3.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)

# Detailed explanations for Polje 1.1, 1.2, 1.3
polje_explanations = [
    ('Polje 1.1', ' se označava znakom «X» ako se radi o ', 'Internom ovlašćenju', ' - ovlašćenje koje daje poreski obveznik (odgovorno lice u pravnom licu, odnosno fizičko lice koje obavlja samostalnu djelatnost) svom zaposlenom licu za zastupanje kod poreskog organa.'),
    ('Polje 1.2', ' se označava znakom «X» ako se radi o ', 'Ugovornom ovlašćenju', ' - ovlašćenje koje daje poreski obveznik (odgovorno lice u pravnom licu, odnosno fizičko lice koje obavlja samostalnu djelatnost) drugom pravnom odnosno fizičkom licu koje je poreski obveznik (npr. knjigovodstvenoj agenciji, advokatskoj kancelariji ili dr.) za zastupanje kod poreskog organa.'),
    ('Polje 1.3', ' se označava znakom «X» ako se radi o ', 'Ostalom ovlašćenju', ' - ovlašćenje koje daje poreski obveznik fizičkom licu koje kao takvo nema registrovanu samostalnu djelatnost. Ovlašćenje mora biti ovjereno kod suda ili notara ili organa lokalne samouprave.'),
]
for polje, text1, bold_inner, text2 in polje_explanations:
    p = doc.add_paragraph()
    run = p.add_run(polje)
    run.font.size = Pt(9)
    run.bold = True
    run2 = p.add_run(text1)
    run2.font.size = Pt(9)
    run3 = p.add_run(bold_inner)
    run3.font.size = Pt(9)
    run3.bold = True
    run3.underline = True
    run4 = p.add_run(text2)
    run4.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)

# Section 2 instruction
p = doc.add_paragraph()
run = p.add_run('2.')
run.font.size = Pt(9)
run.bold = True
run2 = p.add_run('    ')
run2.font.size = Pt(9)
run3 = p.add_run('Podaci o ovlašćenom licu')
run3.font.size = Pt(9)
run3.bold = True
run4 = p.add_run(' – U ovu rubriku upisuju se podaci o ovlašćenom licu: jedinstveni matični broj i prezime i ime. Ukoliko se radi o ugovornom ovlašćenju upisuju se i podaci o PIB-u i nazivu poreskog obveznika sa kojim je zaključen ugovor o zastupanju, kao i zaposlenom licu kod tog poreskog obveznika.')
run4.font.size = Pt(9)
p.paragraph_format.space_after = Pt(6)

# Section 3 instruction
p = doc.add_paragraph()
run = p.add_run('3.')
run.font.size = Pt(9)
run.bold = True
run2 = p.add_run('    ')
run2.font.size = Pt(9)
run3 = p.add_run('Vrsta ovlašćenja')
run3.font.size = Pt(9)
run3.bold = True
run4 = p.add_run(' - U ovu rubriku upisuju se podaci o vrsti ovlašćenja: 3.1 Podnošenje zahtjeva, 3.2 Podnošenje i pregled poreskih prijava; 3.3 Pregled analitičke kartice i 3.4 Pristup e-sandučetu.\nU polje prijava se označava znakom «X» ako se radi o prijavi vrste ovlašćenja, kao i period trajanja ovlašćenja (neograničeno ili ograničeno)\nU polje prestanak se označava znakom «X» ako se radi o prestanku datog ovlašćenja."')
run4.font.size = Pt(9)
p.paragraph_format.space_after = Pt(12)

# Član 5
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Član 5')
run.font.size = Pt(10)
run.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run('Ovaj pravilnik stupa na snagu danom objavljivanja u „Službenom listu Crne Gore", a primjenjivaće se od 1. januara 2026. godine.')
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
run = p.add_run('Broj: 10/1-1-040/25-47246/3\nPodgorica, 29. decembra 2025. godine')
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Ministar,\n')
run.font.size = Pt(10)
run2 = p.add_run('Novica Vuković, ')
run2.font.size = Pt(10)
run2.bold = True
run3 = p.add_run('s.r.')
run3.font.size = Pt(10)

# Save the document
output_path = '/workspace/JPR_2026.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
